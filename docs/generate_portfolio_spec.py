"""
Generate a complete Portfolio Specification & Content Script for Mozn Jamous.
Output: a polished .docx that describes:
1. The strategic positioning of the portfolio
2. The page-by-page structure
3. The exact content/copy to use (English)
4. Design system (colors, typography, spacing)
5. Tech stack recommendations
6. Build phases (8-week plan)
This is a working document that doubles as a brief and a content source.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------- Helpers ----------------
def add_para(doc, text, *, bold=False, size=11, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, font='Calibri', space_after=6,
             italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 28, 1: 20, 2: 15, 3: 12}
    colors = {0: (33, 71, 130), 1: (33, 71, 130), 2: (66, 103, 178), 3: (90, 90, 90)}
    return add_para(doc, text, bold=True, size=sizes.get(level, 11),
                    color=colors.get(level, (0, 0, 0)), space_after=10)


def add_bullet(doc, text, level=0):
    style = 'List Bullet' if level == 0 else f'List Bullet {min(level+1, 3)}'
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(40, 40, 40)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F4F8')
    pPr.append(shd)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'“{text}”')
    run.italic = True
    run.font.name = 'Georgia'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color='214782'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, color=(180, 95, 6)):
    add_para(doc, f"◆ {title}", bold=True, size=12, color=color, space_after=2)
    add_para(doc, body, size=11, color=(60, 60, 60), space_after=10)


def add_section_break(doc):
    add_para(doc, "─" * 60, color=(200, 200, 200),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


# ---------------- Build ----------------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


# ============================ COVER ============================
add_para(doc, "Mozn Jamous", bold=True, size=42, color=(33, 71, 130),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Portfolio Specification & Content Script",
         bold=True, size=22, color=(66, 103, 178),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "A complete blueprint for the personal portfolio website",
         italic=True, size=13, color=(120, 120, 120),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para(doc, "─" * 60, color=(200, 200, 200),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

add_para(doc, "Version 1.0  ·  May 2026", size=11, color=(120, 120, 120),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Prepared for personal use; doubles as a content source for build.",
         size=11, color=(120, 120, 120),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_heading(doc, "What is this document?", level=2)
add_para(doc,
    "This is the single source of truth for the personal portfolio website at "
    "moznjamous.com (proposed domain). It contains the strategic positioning, "
    "the page-by-page structure, the exact English copy to use, the design system, "
    "the tech stack, and the build phases. Treat every section as a brief to a "
    "designer/developer — or as the script you write into the actual site.")

doc.add_page_break()


# ============================ TABLE OF CONTENTS ============================
add_heading(doc, "Contents", level=0)
toc = [
    "1. Strategic Positioning",
    "2. Target Audiences & Primary Calls-to-Action",
    "3. Information Architecture (Site Map)",
    "4. Page-by-Page Specification & Content Script",
    "   4.1  Home",
    "   4.2  About",
    "   4.3  Projects (Index)",
    "   4.4  Project Detail (BloomBelly)",
    "   4.5  Project Detail (CareConnect)",
    "   4.6  Project Detail (Smart Expense Manager)",
    "   4.7  Collaborations",
    "   4.8  Writing / Notes (Future)",
    "   4.9  Contact",
    "5. Design System (Colors, Typography, Spacing, Motion)",
    "6. Tech Stack & Repository Structure",
    "7. Performance, SEO & Accessibility",
    "8. Analytics & Lead Capture",
    "9. Launch Plan (8-Week Phased Build)",
    "10. Maintenance Cadence",
    "11. Appendix: Reusable Snippets & Templates",
]
for line in toc:
    add_para(doc, line, size=11, space_after=2)

doc.add_page_break()


# ============================ 1. STRATEGIC POSITIONING ============================
add_heading(doc, "1. Strategic Positioning", level=0)

add_callout(doc, "One-line positioning",
    "Mozn Jamous — software engineer turning research-grade ideas into "
    "production systems, with a focus on Arabic-first digital health and ERP.")

add_heading(doc, "1.1 Positioning principles", level=2)
add_numbered(doc,
    "Lead with the work, not the credentials. The first thing a visitor sees is "
    "a hero image / video of BloomBelly running on a phone — not a school logo.")
add_numbered(doc,
    "Frame each project around the human problem first, the engineering second. "
    "Reviewers from grants, accelerators, and engineering teams all respond to "
    "the problem statement.")
add_numbered(doc,
    "Show three audiences in one site without splitting them: hiring managers, "
    "grant reviewers, and accelerator scouts each find what they need on the "
    "same page through clear sub-sections.")
add_numbered(doc,
    "Maintain the same look and tone everywhere: portfolio site, GitHub READMEs, "
    "press kit PDFs. Visitors should feel they are inside one ecosystem.")
add_numbered(doc,
    "Be specific and falsifiable. Replace adjectives (\"high-quality\", \"scalable\") "
    "with numbers (\"50+ test cases\", \"4 user roles\", \"17 references\").")

add_heading(doc, "1.2 What this portfolio is NOT", level=2)
add_bullet(doc, "Not a marketing site for the projects — those have their own showcase repos.")
add_bullet(doc, "Not a blog — writing comes later (Section 4.8 is a placeholder).")
add_bullet(doc, "Not a CV in HTML — it is richer than a CV; the downloadable CV is a derivative.")
add_bullet(doc, "Not Arabic-first — the site is English-primary; the projects themselves are Arabic-first.")

doc.add_page_break()


# ============================ 2. TARGET AUDIENCES & CTAs ============================
add_heading(doc, "2. Target Audiences & Primary Calls-to-Action", level=0)

add_table(doc,
    headers=["Audience", "What they need to see", "Primary CTA"],
    rows=[
        ["Hiring managers (Flutter, Odoo, full-stack)",
         "Architecture decisions, code quality signals, real shipped projects",
         "“Download CV” + “Email me”"],
        ["Grant reviewers (digital health)",
         "Problem framing, evidence base, theory of change, team credibility",
         "“Read the BloomBelly case study” + “Download press kit”"],
        ["Accelerator scouts (MENA, healthtech, fintech)",
         "Business model, market sizing, traction, founder-market fit",
         "“Book a 15-min call” + “Download press kit”"],
        ["Academic supervisors / researchers",
         "Methodology, references, supervisors, thesis context",
         "“Citation file” + “Email me”"],
        ["Future collaborators (Shahd, shaha123s-type partners)",
         "Honest credit, clear contribution split, joy of collaboration",
         "“See collaborations” + “Email me”"],
    ],
    col_widths=[5.5, 6, 5])

add_heading(doc, "2.1 Friction-removal rules", level=2)
add_bullet(doc, "No “Contact me” form on the home page — just an email and a LinkedIn link.")
add_bullet(doc, "No newsletter pop-up. Ever.")
add_bullet(doc, "No language switcher in V1 — the site is English-only.")
add_bullet(doc, "No cookie banner unless analytics requires one; if so, single-line, dismiss-able.")
add_bullet(doc, "No more than two CTAs visible above the fold on any page.")

doc.add_page_break()


# ============================ 3. INFORMATION ARCHITECTURE ============================
add_heading(doc, "3. Information Architecture (Site Map)", level=0)

add_code(doc, """moznjamous.com
├── /                         Home
├── /about                    About + story + skills + education
├── /projects                 Projects index (3 featured cards)
│   ├── /projects/bloombelly       Graduation thesis, deepest case study
│   ├── /projects/careconnect      Three-app marketplace
│   └── /projects/smart-expense    Odoo module (open source)
├── /collaborations           Production codebases I contribute to
├── /writing                  (V2) Long-form essays and ADRs
├── /press-kit                Download links for press materials
├── /cv                       Downloadable PDF CV
└── /contact                  Email + LinkedIn + scheduling link""")

add_heading(doc, "3.1 Navigation rules", level=2)
add_bullet(doc, "Top nav (max 5 items): Projects · Collaborations · About · Contact · CV (button).")
add_bullet(doc, "Footer mirrors the nav and adds GitHub, LinkedIn, email, press-kit, sitemap.")
add_bullet(doc, "Project detail pages are always reachable by URL even if the index changes.")
add_bullet(doc, "Every page has “What's next?” links at the bottom to guide the journey.")

doc.add_page_break()


# ============================ 4. PAGE-BY-PAGE SPEC ============================
add_heading(doc, "4. Page-by-Page Specification & Content Script", level=0)

# 4.1 Home
add_heading(doc, "4.1 Home (/)", level=2)
add_para(doc, "Goal: convince the visitor within 6 seconds that this is a serious engineer working on real problems.", italic=True, color=(120, 120, 120))

add_heading(doc, "Hero section (above the fold)", level=3)
add_para(doc, "Layout: 60/40 split — left side text, right side a phone mockup of BloomBelly.")
add_para(doc, "Copy (exact text to use):", bold=True)
add_quote(doc, "Hi, I'm Mozn Jamous. I build production systems for problems "
              "that matter — Arabic-first digital health, Odoo ERP, and "
              "AI-powered mobile apps. Currently graduating from Al-Sham "
              "Private University, Damascus.")
add_para(doc, "Two buttons under the hero copy:")
add_bullet(doc, "Primary: “See my work” → scrolls to projects")
add_bullet(doc, "Secondary: “Download CV (PDF)”")

add_heading(doc, "Featured projects section", level=3)
add_para(doc, "Three cards in a row. Each card has: thumbnail, project name, one-line description, tech tags, link.")
add_para(doc, "Order matters — always BloomBelly first, then CareConnect, then Smart Expense Manager.")

add_heading(doc, "Selected collaborations section", level=3)
add_para(doc, "Short text + three logos/cards for the production codebases. Each links to /collaborations.")

add_heading(doc, "Quick-fact strip", level=3)
add_para(doc, "Horizontal strip with four stats:")
add_table(doc,
    headers=["Stat", "Value", "Tooltip on hover"],
    rows=[
        ["Projects shipped", "11+", "Across solo and team contexts"],
        ["Lines of production code", "20M+", "Aggregate across owned + collaborative repos"],
        ["User roles designed", "30+", "Across BloomBelly, CareConnect, takhrjy, Mademoiselle"],
        ["Languages I build with", "8", "Dart, Python, TypeScript, Vue, JS, SQL, PL/pgSQL, Bash"],
    ],
    col_widths=[5, 3, 8])

add_heading(doc, "Closing band", level=3)
add_para(doc, "One line of copy + email + LinkedIn:")
add_quote(doc, "If you're working on something I can help with — a digital health grant, "
              "an Odoo build, an Arabic-first product — I'd love to hear about it.")

# 4.2 About
doc.add_page_break()
add_heading(doc, "4.2 About (/about)", level=2)
add_para(doc, "Goal: build trust through narrative — who I am, why I work on what I work on.", italic=True, color=(120, 120, 120))

add_heading(doc, "Opening", level=3)
add_quote(doc, "I'm a final-year Informatics Engineering student at Al-Sham Private "
              "University in Damascus, graduating in 2026. The work I'm proudest of "
              "comes from a single conviction: software is most valuable when it shows "
              "up where the world is hardest — Arabic-language healthcare, "
              "small-business ERP, the daily struggle of working mothers.")

add_heading(doc, "Sections", level=3)
add_bullet(doc, "How I got here — short paragraph about Damascus, ASPU, why software engineering.")
add_bullet(doc, "What I build with — a stack overview with rationale (NOT a long list).")
add_bullet(doc, "How I work — collaboration principles, communication preferences, response times.")
add_bullet(doc, "What I'm looking for — the explicit asks: jobs, grants, accelerators, partnerships.")
add_bullet(doc, "Education — ASPU details, supervisors, thesis.")
add_bullet(doc, "Outside of work — one short, human paragraph.")

add_heading(doc, "What NOT to include", level=3)
add_bullet(doc, "A 30-skill cloud. Pick 8.")
add_bullet(doc, "A long timeline of every course. Just the relevant ones.")
add_bullet(doc, "A list of every certification. Pick the strongest 3 if any.")

# 4.3 Projects index
doc.add_page_break()
add_heading(doc, "4.3 Projects Index (/projects)", level=2)
add_para(doc, "Goal: let visitors pick the project that's most relevant to them.", italic=True, color=(120, 120, 120))

add_para(doc, "Layout: three large project cards stacked vertically, then a 'Collaborations' teaser at the bottom.")
add_para(doc, "Each card contains:", bold=True)
add_bullet(doc, "Project name + emoji icon")
add_bullet(doc, "One-line description")
add_bullet(doc, "Status badge (Production, Graduation Thesis, Archived, etc.)")
add_bullet(doc, "Three tech tags")
add_bullet(doc, "Two CTAs: “Read case study” and “View on GitHub”")
add_bullet(doc, "Hero image (Figma frame or app screenshot)")

# 4.4 BloomBelly detail
doc.add_page_break()
add_heading(doc, "4.4 Project Detail: BloomBelly (/projects/bloombelly)", level=2)
add_para(doc, "Goal: convince a grant reviewer, employer, or accelerator that this is real and grounded.", italic=True, color=(120, 120, 120))

add_heading(doc, "Page structure", level=3)
add_numbered(doc, "Hero: title, tagline, status badge, three tech tags, hero image")
add_numbered(doc, "TL;DR: 3-sentence summary anyone can absorb in 10 seconds")
add_numbered(doc, "The Problem: maternal/child health gap in Arabic-speaking regions, with statistics")
add_numbered(doc, "The Solution: features list with icons, organized by user role")
add_numbered(doc, "Architecture: high-level diagram (Flutter — Flask — Gemini/LoRA/RF — Supabase)")
add_numbered(doc, "Evidence: research references table with 7 sources")
add_numbered(doc, "Engineering decisions: 3-4 of the most interesting ADRs with one-paragraph each")
add_numbered(doc, "Team: Mozn + Shahd photos (if comfortable) + supervisors")
add_numbered(doc, "Outcomes: what shipped, what's in the pipeline, what we learned")
add_numbered(doc, "Links bar: showcase repo, Figma, citation, press kit, contact")

add_heading(doc, "Hero copy (exact)", level=3)
add_quote(doc, "BloomBelly — An Arabic-first maternal and child health companion. "
              "Built as a graduation thesis with Shahd Bureghsh under "
              "Dr. Afaf Al-Shalabi and Eng. Rahaf Abdul Qader at "
              "Al-Sham Private University.")

add_heading(doc, "TL;DR copy (exact)", level=3)
add_quote(doc, "BloomBelly combines three specialized AI components — Google "
              "Gemini for medical image analysis, a LoRA-fine-tuned transformer for "
              "the medical chatbot, and a Random Forest classifier for nutrition "
              "guidance — behind a Python Flask backend and a Flutter mobile "
              "app, with a doctor-administered wallet system that enables clinic "
              "distribution.")

# 4.5 CareConnect detail
doc.add_page_break()
add_heading(doc, "4.5 Project Detail: CareConnect (/projects/careconnect)", level=2)
add_para(doc, "Goal: showcase system-design skill and partnership working.", italic=True, color=(120, 120, 120))

add_heading(doc, "Same structure as BloomBelly, with these differences", level=3)
add_bullet(doc, "The hero emphasizes the three-app ecosystem, not the AI")
add_bullet(doc, "The architecture diagram is the centerpiece (3 apps + Supabase + Maps/WhatsApp)")
add_bullet(doc, "The 'Engineering decisions' section pulls 3 ADRs from the showcase repo: “Three apps over one”, “Supabase over Firebase”, “RBAC at the DB layer”")
add_bullet(doc, "The team section gives the Babysitter App co-developer (@shaha123s) prominent credit")
add_bullet(doc, "The links bar features the Figma prototype prominently")

# 4.6 Smart Expense detail
doc.add_page_break()
add_heading(doc, "4.6 Project Detail: Smart Expense Manager (/projects/smart-expense)", level=2)
add_para(doc, "Goal: prove I can ship clean, open-source backend code in a niche stack (Odoo).", italic=True, color=(120, 120, 120))

add_heading(doc, "Structure", level=3)
add_numbered(doc, "Hero: Odoo 19 module emphasized; license badge (LGPL-3.0)")
add_numbered(doc, "The Problem: paper-based expense workflows in small businesses")
add_numbered(doc, "The Solution: mobile receipts, email approval flow, automated journal entries")
add_numbered(doc, "Workflow diagram: Employee → Manager email → One-click approve → Finance posting")
add_numbered(doc, "Why open source: educational, builds Odoo-developer reputation, low-risk IP")
add_numbered(doc, "How to use it: link to GitHub README installation steps")
add_numbered(doc, "Tech depth: a few snippets (anonymized) showing token-based auth pattern")

# 4.7 Collaborations
doc.add_page_break()
add_heading(doc, "4.7 Collaborations (/collaborations)", level=2)
add_para(doc, "Goal: show I work well with others and contribute to production code I don't own.", italic=True, color=(120, 120, 120))

add_heading(doc, "Opening copy", level=3)
add_quote(doc, "Some of my best work lives in repos I don't own. These are "
              "production codebases I contribute to as part of partner teams. "
              "Where possible, I link to the partner's GitHub so you can verify "
              "the collaboration is real.")

add_heading(doc, "Three collaboration cards", level=3)
add_para(doc, "Each card includes: partner name, project, my contribution areas, dates, status.")
add_para(doc, "Cards to feature:")
add_bullet(doc, "Mademoiselle (V1 + V2) — with @santateammedia")
add_bullet(doc, "eda — with @xvtu2003")
add_bullet(doc, "Babysitter Provider App — with @shaha123s (sister of CareConnect)")

add_heading(doc, "Honesty principle", level=3)
add_para(doc, "Be explicit about what I did vs. what the team did. Use phrases like:")
add_bullet(doc, "“Contributed to the backend RPCs” — not “Built”")
add_bullet(doc, "“Co-designed the booking flow” — not “Designed”")
add_bullet(doc, "“Paired on the auth refactor” — not “Refactored auth”")

# 4.8 Writing (future)
doc.add_page_break()
add_heading(doc, "4.8 Writing / Notes (Future) (/writing)", level=2)
add_para(doc, "Goal (future): demonstrate thinking through long-form writing.", italic=True, color=(120, 120, 120))
add_para(doc, "In V1: this page does not exist. Top nav shows it as “Writing (soon)” only if scoping a V2 in the next 6 months.")
add_para(doc, "When activated, candidate post topics:")
add_bullet(doc, "“How we fine-tuned a transformer for Arabic medical Q&A on a graduation budget”")
add_bullet(doc, "“Why we chose Supabase + Flask over Firebase for a healthcare app”")
add_bullet(doc, "“Designing a three-app childcare marketplace with Supabase RLS”")
add_bullet(doc, "“Token-based email approvals: a security pattern for low-friction workflows”")
add_bullet(doc, "“What I learned writing a 50-page SRS as an undergraduate”")

# 4.9 Contact
doc.add_page_break()
add_heading(doc, "4.9 Contact (/contact)", level=2)
add_para(doc, "Goal: convert interested visitors into a first email or call.", italic=True, color=(120, 120, 120))

add_heading(doc, "Contact page content", level=3)
add_quote(doc, "If you're working on something I can help with — hiring for a "
              "Flutter, Odoo, or full-stack role; running a digital health grant; "
              "investing in MENA health tech; or just curious about something on "
              "this site — the fastest way to reach me is email.")
add_para(doc, "Three blocks:")
add_bullet(doc, "Email: moznjamous9@gmail.com (with “Copy” button)")
add_bullet(doc, "LinkedIn: linkedin.com/in/mozn-jamous (with profile preview)")
add_bullet(doc, "Calendar: optional Cal.com / Calendly link for a 15-min intro call")

add_heading(doc, "What to say in the email subject line", level=3)
add_para(doc, "Suggested template snippets the page can show as examples:")
add_bullet(doc, "“[Hiring] Flutter role at [company]”")
add_bullet(doc, "“[Grant] [program name] — interested in BloomBelly”")
add_bullet(doc, "“[Partnership] [project area]”")
add_bullet(doc, "“[Question] about [project]”")

doc.add_page_break()


# ============================ 5. DESIGN SYSTEM ============================
add_heading(doc, "5. Design System", level=0)

add_heading(doc, "5.1 Color tokens", level=2)
add_table(doc,
    headers=["Token", "Hex", "Use"],
    rows=[
        ["--ink", "#0F172A", "Body text, headings"],
        ["--ink-muted", "#475569", "Captions, secondary text"],
        ["--paper", "#FAFAFA", "Page background"],
        ["--surface", "#FFFFFF", "Cards, sections"],
        ["--border", "#E2E8F0", "Dividers, subtle outlines"],
        ["--brand-primary", "#214782", "Headings, primary buttons (BloomBelly navy)"],
        ["--brand-accent", "#4267B2", "Links, hover states"],
        ["--brand-warm", "#B45F06", "Highlights, badges (use sparingly, <3%)"],
        ["--success", "#16A34A", "Confirmations"],
        ["--danger", "#DC2626", "Errors, warnings"],
    ],
    col_widths=[4, 3, 7])

add_heading(doc, "5.2 Typography", level=2)
add_para(doc, "Pair one display serif with one neutral sans-serif:")
add_bullet(doc, "Display: “Cormorant Garamond” (for hero text, project titles)")
add_bullet(doc, "Body: “Inter” (for everything else)")
add_bullet(doc, "Monospace: “JetBrains Mono” (for code blocks)")
add_para(doc, "Type scale (rem):")
add_code(doc, """--text-xs:   0.75rem
--text-sm:   0.875rem
--text-base: 1rem
--text-lg:   1.125rem
--text-xl:   1.25rem
--text-2xl:  1.5rem
--text-3xl:  1.875rem
--text-4xl:  2.25rem
--text-5xl:  3rem
--text-6xl:  3.75rem  (hero)""")

add_heading(doc, "5.3 Spacing scale", level=2)
add_para(doc, "Tailwind-style 4-pixel base; common values you will use:")
add_code(doc, "0.25rem  0.5rem  0.75rem  1rem  1.5rem  2rem  3rem  4rem  6rem  8rem")

add_heading(doc, "5.4 Layout grid", level=2)
add_bullet(doc, "Max content width: 1180 px")
add_bullet(doc, "Mobile padding: 1.25rem")
add_bullet(doc, "Desktop padding: 2rem")
add_bullet(doc, "Section vertical rhythm: 6rem between major sections; 3rem inside")

add_heading(doc, "5.5 Motion", level=2)
add_bullet(doc, "Durations: 140ms (micro), 240ms (default), 360ms (page-level)")
add_bullet(doc, "Easing: cubic-bezier(0.4, 0, 0.2, 1)")
add_bullet(doc, "Respect prefers-reduced-motion: reduce")
add_bullet(doc, "No parallax. No auto-playing video. No carousel with auto-rotate.")

add_heading(doc, "5.6 Imagery rules", level=2)
add_bullet(doc, "Phone mockups for mobile-app screens; never raw screenshots without a frame")
add_bullet(doc, "Architecture diagrams in PNG export from Excalidraw / Figma; never as live Mermaid (Mermaid is fine for GitHub but not for the polished site)")
add_bullet(doc, "Reuse the brand navy/accent in diagrams so they feel native to the page")
add_bullet(doc, "Compress with WebP + AVIF fallbacks")

doc.add_page_break()


# ============================ 6. TECH STACK ============================
add_heading(doc, "6. Tech Stack & Repository Structure", level=0)

add_heading(doc, "6.1 The recommended stack", level=2)
add_table(doc,
    headers=["Concern", "Choice", "Why"],
    rows=[
        ["Framework", "Next.js 14 (App Router)", "SSG + ISR; perfect for content-first sites; mature ecosystem"],
        ["Language", "TypeScript", "Type safety; matches the seriousness of the projects"],
        ["Styling", "Tailwind CSS", "Token-friendly; aligns with the design system"],
        ["Components", "shadcn/ui (selectively)", "High-quality unstyled primitives; copy-paste, no lock-in"],
        ["Content", "MDX", "Lets case studies mix prose, code, and components"],
        ["Icons", "lucide-react", "Consistent, clean, free"],
        ["Hosting", "Vercel", "Best-in-class for Next.js; free for personal use"],
        ["Domain", "moznjamous.com (suggested)", "Personal-brand domain; ~$10/year"],
        ["Analytics", "Plausible or Vercel Analytics", "Privacy-respecting; no cookie banner needed"],
        ["Forms (future)", "Resend + a simple API route", "Lightweight, no third-party form service"],
    ],
    col_widths=[3.5, 4.5, 7])

add_heading(doc, "6.2 Repository structure", level=2)
add_code(doc, """personal-portfolio/
├── app/
│   ├── (marketing)/
│   │   ├── page.tsx                 # Home
│   │   ├── about/page.tsx
│   │   ├── collaborations/page.tsx
│   │   ├── contact/page.tsx
│   │   └── press-kit/page.tsx
│   ├── projects/
│   │   ├── page.tsx                 # Index
│   │   ├── bloombelly/page.mdx
│   │   ├── careconnect/page.mdx
│   │   └── smart-expense/page.mdx
│   └── layout.tsx
├── components/
│   ├── hero.tsx
│   ├── project-card.tsx
│   ├── stat-strip.tsx
│   ├── collaboration-card.tsx
│   ├── nav.tsx
│   └── footer.tsx
├── content/
│   └── cv.pdf                       # Static PDF
├── public/
│   ├── og/                          # OG images per page
│   ├── mockups/                     # Phone frames
│   ├── diagrams/                    # Architecture PNGs
│   └── press-kit/                   # Downloadables
├── styles/
│   └── globals.css                  # Tokens + Tailwind
├── next.config.mjs
└── package.json""")

doc.add_page_break()


# ============================ 7. PERFORMANCE, SEO, A11Y ============================
add_heading(doc, "7. Performance, SEO & Accessibility", level=0)

add_heading(doc, "7.1 Performance targets", level=2)
add_table(doc,
    headers=["Metric", "Target", "How"],
    rows=[
        ["Lighthouse Performance", "≥ 95", "Static export where possible; image optimization"],
        ["Largest Contentful Paint", "< 1.5s on 4G", "Above-the-fold hero is a static image"],
        ["Cumulative Layout Shift", "< 0.05", "Reserve image dimensions"],
        ["Total bundle (gzipped)", "< 100 KB", "Tree-shake; no chart library on home"],
        ["Time to first byte", "< 200ms", "Vercel edge runtime where applicable"],
    ],
    col_widths=[5, 3, 6])

add_heading(doc, "7.2 SEO essentials", level=2)
add_bullet(doc, "Unique <title> and <meta name=\"description\"> on every page")
add_bullet(doc, "Open Graph image (1200×630) for each project page; brand-aligned")
add_bullet(doc, "Schema.org/Person on /about; CreativeWork on each project page")
add_bullet(doc, "Sitemap.xml + robots.txt; both submitted to Google Search Console")
add_bullet(doc, "Canonical URLs; no duplicate content")

add_heading(doc, "7.3 Accessibility (WCAG 2.1 AA target)", level=2)
add_bullet(doc, "Keyboard-navigable nav and footer; visible focus rings")
add_bullet(doc, "Color contrast ≥ 4.5:1 for body text; ≥ 3:1 for large text")
add_bullet(doc, "Semantic HTML (header, nav, main, article, section, footer)")
add_bullet(doc, "Alt text on every image; aria-labels on icon-only buttons")
add_bullet(doc, "Skip-to-content link as first focusable element")
add_bullet(doc, "Forms (when added): labels, error messages, validation feedback")

doc.add_page_break()


# ============================ 8. ANALYTICS & LEAD CAPTURE ============================
add_heading(doc, "8. Analytics & Lead Capture", level=0)

add_heading(doc, "8.1 What to track", level=2)
add_bullet(doc, "Page views per page (with referrer)")
add_bullet(doc, "CV download clicks")
add_bullet(doc, "Email “copy” button clicks")
add_bullet(doc, "Press-kit downloads (per project)")
add_bullet(doc, "Outbound clicks to GitHub repos")
add_bullet(doc, "Outbound clicks to Figma prototypes")

add_heading(doc, "8.2 What NOT to track", level=2)
add_bullet(doc, "Scroll depth (noisy; rarely actionable)")
add_bullet(doc, "Mouse movement / heatmaps (creepy; legally risky)")
add_bullet(doc, "Anything PII without explicit consent")

add_heading(doc, "8.3 Lead-capture flows", level=2)
add_callout(doc, "V1 — the email-first flow",
    "Visitor reads, becomes interested, clicks the email link in nav/footer. "
    "I get an email. Done. No form, no popup, no third-party service.")
add_callout(doc, "V2 — the calendar flow (optional)",
    "Add a Cal.com or Calendly link on /contact for a 15-minute intro call. "
    "Useful once email volume justifies it.",
    color=(0, 120, 60))

doc.add_page_break()


# ============================ 9. LAUNCH PLAN ============================
add_heading(doc, "9. Launch Plan (8-Week Phased Build)", level=0)

add_heading(doc, "Week 1 — Foundation", level=2)
add_bullet(doc, "Buy moznjamous.com")
add_bullet(doc, "Create new Next.js project; deploy a blank build to Vercel")
add_bullet(doc, "Implement nav, footer, design tokens; ship a placeholder home page")
add_bullet(doc, "Set up GitHub Actions + Vercel preview deployments")

add_heading(doc, "Week 2 — Home + About", level=2)
add_bullet(doc, "Build Home (hero, featured projects, stat strip, closing band)")
add_bullet(doc, "Build About (story, stack, looking-for, education)")
add_bullet(doc, "Take and place phone mockups")

add_heading(doc, "Week 3 — Projects Index + BloomBelly detail", level=2)
add_bullet(doc, "Build Projects index with the three cards")
add_bullet(doc, "Build BloomBelly detail page; export the architecture diagram as PNG")
add_bullet(doc, "Wire links to the showcase repo and press kit")

add_heading(doc, "Week 4 — CareConnect + Smart Expense detail pages", level=2)
add_bullet(doc, "Build both project detail pages; mirror BloomBelly's structure")
add_bullet(doc, "Embed Figma prototype on CareConnect page")

add_heading(doc, "Week 5 — Collaborations + Contact", level=2)
add_bullet(doc, "Build /collaborations with three cards")
add_bullet(doc, "Build /contact with email-copy button + LinkedIn block")
add_bullet(doc, "Generate clean OG images for every page")

add_heading(doc, "Week 6 — Press kit + CV", level=2)
add_bullet(doc, "Build /press-kit page; upload PDF press kits from showcase repos")
add_bullet(doc, "Generate final CV PDF; make /cv route serve it")
add_bullet(doc, "Add structured data (Schema.org) per page")

add_heading(doc, "Week 7 — Polish + Performance", level=2)
add_bullet(doc, "Run Lighthouse; fix anything < 95")
add_bullet(doc, "Optimize images (WebP/AVIF)")
add_bullet(doc, "Accessibility audit; keyboard navigation review")
add_bullet(doc, "Cross-browser testing")

add_heading(doc, "Week 8 — Launch", level=2)
add_bullet(doc, "Submit sitemap to Google Search Console")
add_bullet(doc, "Set up Plausible / Vercel Analytics")
add_bullet(doc, "Announce on LinkedIn with a tasteful post")
add_bullet(doc, "Email a handful of friends/mentors for feedback")
add_bullet(doc, "Push v1.0 tag")

doc.add_page_break()


# ============================ 10. MAINTENANCE CADENCE ============================
add_heading(doc, "10. Maintenance Cadence", level=0)

add_table(doc,
    headers=["Cadence", "Task"],
    rows=[
        ["Weekly", "Review analytics; capture any anomalies; respond to incoming email"],
        ["Monthly", "Update project status badges if anything shipped or paused"],
        ["Quarterly", "Refresh stats on the home page; revisit the “looking for” section"],
        ["Twice a year", "Add a new project or major update; rotate featured collaborations"],
        ["Annually", "Major design review; consider V2 of the site"],
    ],
    col_widths=[4, 11])

doc.add_page_break()


# ============================ 11. APPENDIX ============================
add_heading(doc, "11. Appendix: Reusable Snippets & Templates", level=0)

add_heading(doc, "11.1 Project-card component spec", level=2)
add_code(doc, """type ProjectCardProps = {
  title: string;
  emoji: string;
  oneLiner: string;
  status: 'Production' | 'Graduation Thesis' | 'Open Source' | 'Archived';
  tags: string[];                  // max 3
  heroImage: string;               // 16:9
  href: string;                    // /projects/[slug]
  externalLinks: {
    github?: string;
    figma?: string;
    pressKit?: string;
  };
};""")

add_heading(doc, "11.2 Standard email signature", level=2)
add_code(doc, """Mozn Jamous
Informatics Engineer (graduating 2026, ASPU)
moznjamous9@gmail.com  ·  github.com/Mozn-jamous  ·  linkedin.com/in/mozn-jamous""")

add_heading(doc, "11.3 Standard LinkedIn-post template for project launches", level=2)
add_code(doc, """[1-sentence problem]

[1-sentence solution]

[1-sentence what was hardest / what we learned]

🔗 [showcase URL]

Built with [@collaborator if any]. Supervised by [supervisor if any].
""")

add_heading(doc, "11.4 Hero “What I do” one-liner (reusable in many places)", level=2)
add_quote(doc, "I build production systems for problems that matter — "
              "Arabic-first digital health, Odoo ERP, and AI-powered mobile apps.")

add_heading(doc, "11.5 Closing principle", level=2)
add_quote(doc, "A portfolio is not a museum of past work. It is an invitation. "
              "Everything on this site should be designed to make a "
              "specific reader feel: 'I want to work with this person.'")

add_para(doc, "")
add_section_break(doc)
add_para(doc, "End of document. Version 1.0 — May 2026.",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=(120, 120, 120))

# Save
output = r"c:\Users\mesho\OneDrive\Desktop\portfolio-website\Portfolio_Specification_v1.docx"
doc.save(output)
print(f"Saved: {output}")
